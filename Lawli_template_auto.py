#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Sep 13 16:21:45 2018

@author: haibinli
"""
import docx
from docx import Document
import pandas as pd
import os
import codecs
import re
import nltk

import argparse
import datetime
from dateutil.parser import parse

from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time,datetime
import shutil
import configparser
from datetime import datetime

import sys


path_q = Document('/Users/haibinli/Lawli/Code/Game Publishing Agreement.Template.docx')
path_a = Document('/Users/haibinli/Lawli/Code/Questions for Demo 2.docx')

file_q = '/Users/haibinli/Lawli/Code/Game Publishing Agreement.Template.docx'
file_a = '/Users/haibinli/Lawli/Code/Questions for Demo 2.docx'


print(path_q.paragraphs)
print(path_a.paragraphs)


for q in path_q.paragraphs:
    print(q.text)

for a in path_a.paragraphs:
    print(a.text)

path_a.styles



def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)




full_text_q = getText(file_q)
full_text_a = getText(file_a)



    
def cleanhtml(raw_html):
  cleanr = re.compile('<.*?>')
  cleantext = re.sub(cleanr, '', raw_html)
  return cleantext


# full_text2_clean = cleanhtml(full_text2)


####################### separate document to different sections #####################

doc_a_sents = full_text_a.split("\n")


doc_question = []
doc_answer= []

QA_dict ={}



for i in range(len(doc_question)):
    
    clean_title = re.match(r'(.*)\(.*\)', doc_question[i].lower()) 
    if clean_title:
        QA_sum['Question'][i] = clean_title.group(1).rstrip() 
    else: QA_sum['Question'][i] = doc_question[i]
    
    QA_sum['Mark'][i] = doc_question[i][doc_question[i].find("("):doc_question[i].find(")")+1]
    
         
    try:
        input_arg = sys.argv
        QA_sum['Answer'][i] = input_arg[i]
    
    except IndexError:
        QA_sum['Answer'][i] = input( QA_sum['Question'][i])
        

tempFile = open(file_q, 'r+' )   
    
doc = Document(file_q)      













